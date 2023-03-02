import os
import docx
import PyPDF2
from bs4 import BeautifulSoup

my_directory = os.path.abspath(os.getcwd())

# Función para leer XML (HTML)


def BS(markup):
    return BeautifulSoup(markup, "lxml")

# Write HTMl documents after OpenAI response


def write_html_history(xml, model, section):

    # Count amount of files in ./history folder to asign a version number in file name
    count = 0
    # Iterate directory
    for path in os.listdir(r'history/'):
        # check if current path is a file
        if os.path.isfile(os.path.join(r'history/', path)):
            count += 1

    history_file_name = f'history/v{count}-Syllabus_{section}_{model}.html'
    with open(history_file_name, "w", encoding="UTF-8") as file:
        file.write(str(xml))


# Read the
def read_syllabus_HTML():
    history_files = []
    for path in os.listdir(r'history/'):
        # check if current path is a file
        if os.path.isfile(os.path.join(r'history/', path)):
            history_files.append(path)
    # Open the last created file in syllabus history folder otherwise the original syllabus
    file_to_read = os.path.join(r'history/', history_files[-1]) if len(
        history_files) > 0 else r"syllabus/originalSyllabus.html"
    print(f"Using the syllabus from: {file_to_read}")
    with open(file_to_read, encoding="UTF-8") as html_doc:
        return BS(html_doc)

# Get text from Word document


def read_docx(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text_from_file)
    return '\n'.join(full_text)

# Get text from PDF


def read_pdf(filename):
    pdf_file = open(filename, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    full_text = []
    for page in range(len(pdf_reader.pages)):
        page_obj = pdf_reader.pages[page]
        full_text.append(page_obj.extract_text())
    return '\n'.join(full_text).strip('\n')

# Write word documents


def write_docx(HTML):
    document = docx.Document()
    document.add_heading(HTML.find("title").text)
    document.add_heading("Justificación")
    document.add_paragraph(HTML.find("justification").text)
    document.add_heading("Objetivo")
    document.add_paragraph(HTML.find("objective").text)
    document.save(r"final/final.docx")

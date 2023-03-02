# -*- coding: utf-8 -*-

import docx
from docx import Document
import PyPDF2
import openai
import os
from datetime import datetime
from bs4 import BeautifulSoup
import  tkinter as tk

# Scripts
from get_inputs import capture_openAI_model, capture_syllabus_selected_section

# Librería para ejecutar GUI
# from tkinter import filedialog
# root = tk.Tk()
# root.withdraw()

# Función para leer XML (HTML)


def BS(markup):
    return BeautifulSoup(markup, "lxml")


my_directory = os.path.dirname(__file__)

# Set OpenAI API key
openai.api_key = "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"

# Character's limit for openAI API
openai_maximum_content_length = 2049

# Trim text for OpenAI maximum context length


def trim_text(text_from_file):
    correct_text_length = openai_maximum_content_length - len(openai_prompt)
    preprocess_text = text_from_file.replace("\n", " ")
    trims = [(preprocess_text[i:i + correct_text_length])
             for i in range(0, len(text_from_file), correct_text_length)]
    return trims[0]

# Function to get text from Word document


def read_docx(filename):
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text_from_file)
    return '\n'.join(full_text)

# Function to get text from PDF


def read_pdf(filename):
    pdf_file = open(filename, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    full_text = []
    for page in range(len(pdf_reader.pages)):
        page_obj = pdf_reader.pages[page]
        full_text.append(page_obj.extract_text())
    return '\n'.join(full_text).strip('\n')

# Function to get text from file


def get_text_from_file(file_to_read):
    # Get file extension
    file_extension = file_to_read.split('.')[-1].lower()

    # Read text from file based on extension
    if file_extension == 'docx' or file_extension == 'doc':
        extracted_text = read_docx(file_to_read)
        print("\n__Word document read__\n")
    elif file_extension == 'pdf':
        extracted_text = read_pdf(file_to_read)
        print("Reading PDF document")
    else:
        print('Unsupported file format')
        extracted_text = None
    return extracted_text

# Function to get openAI response


def call_openAI_API(prompt, model, text, section):

    print(
        f"\n______________\n\n Calling OpenAI with the prompt:\n -{prompt}- \n\n Using the model: '{model}' \n\n With the section: '{section}' \n\n______________\n\n Waiting for OpenAI response...\n")

    composedPrompt = f'{prompt} \n text: """{text}"""'

    # buenas prácticas para crear propmts
    # https://help.openai.com/en/articles/6654000-best-practices-for-prompt-engineering-with-openai-api

    if model == "edition":
        completions = openai.Edit.create(
            engine="text-davinci-edit-001",
            input=text,
            instruction=prompt,
            n=1,
            temperature=0.7,
        )

    if model == "completion":
        completions = openai.Completion.create(
            engine="text-davinci-003",
            prompt=composedPrompt,
            max_tokens=1024,
            n=1,
            stop=None,
            temperature=0.7,
        )

    # Get openAI response text
    response = completions.choices[0].text

    print(f'\n______________\n\nOpenAI responded with: \n\n {response}\n\n______________\n')

    return response

# Function to write word document with given prompt and response


def write_docx(prompt, response):
    document = Document()
    document.add_heading("Documento generado automáticamente con el prompt:")
    document.add_paragraph(prompt)
    document.add_heading("Respuesta de OpenAI:")
    document.add_paragraph(response)
    document.save('OpenAI response.docx')


 # 5. Reconstruir el documento y guardarlo
    #   - Duplicar el archivo original (agregar fecha y hora) justification_edit_12-01-23_01:12.html

def write_html_history(xml, model, section) :
    history_file_name = f'history/Syllabus_{datetime.now().strftime("%d-%m-%Y__%H-%M-%S")}_{section}_{model}.html'
    with open(history_file_name, "w", encoding="UTF-8") as file:
        file.write(str(xml))
    return


def read_syllabus_HTML():
    with open(os.path.join(my_directory, "syllabus/FormatedSyllabus.html"), encoding="UTF-8") as html_doc:
        return BS(html_doc)


if __name__ == "__main__":

    print("\n\n~~___Running___~~\n\n")

    # 1. Capturar y almacenar la sección del syllabus que se desea usar
    syllabus_selected_section = capture_syllabus_selected_section()

    # 2. Leer el HTML del syllabus y almacenarlo en una variable
    syllabus_virtual_document = read_syllabus_HTML()

    # 3. Capturar que modelo de OpenAI (Copletion, Edition) desea usar
    openai_model = capture_openAI_model()

    # 4. Capturar las instrrucciones para la IA
    openai_prompt = input(
        "\nType something to do with the text: \nR:")

    # 5. Extraer la sección del texto seleccionado del syllabus
    openai_text = syllabus_virtual_document.find(
        syllabus_selected_section).text

    # # 6. Llamar la API de OpenAI
    # openAI_response = call_openAI_API(
    #     openai_prompt, openai_model, openai_text, syllabus_selected_section)
    
    # 7. reconstructiong the virtual syllabus with the API response
    syllabus_virtual_document.find(syllabus_selected_section).string = "cambio"

    write_html_history(syllabus_virtual_document, openai_model, syllabus_selected_section)
   
